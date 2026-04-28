from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = r"c:\Users\samij\Desktop\uppgifter\alyaa"

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.size = Pt(12)

def spacer(doc):
    doc.add_paragraph("")

# ─────────────────────────────────────────────
# TASK 1 — VISITKORT ANALYS
# ─────────────────────────────────────────────
doc1 = Document()

heading(doc1, "Analys och motivering av visitkort", level=1)
body(doc1, "Alyaa\nGrafisk kommunikation")
spacer(doc1)

heading(doc1, "1. Syfte och målgrupp", level=2)
body(doc1,
    "Syftet med mitt visitkort är att presentera mig som frilansfotograf och göra det lätt för "
    "potentiella kunder att kontakta mig. Visitkortet ska ge ett professionellt intryck men ändå "
    "kännas personligt och kreativt, ungefär som en liten del av min stil.\n\n"
    "Målgruppen är vuxna och unga vuxna, typ 20–45 år, som är intresserade av att anlita en "
    "fotograf för porträtt, event eller produktfotografering. Dom är vana vid att se snygg design "
    "och förväntar sig att visitkortet matchar den fotografens känsla och estetik.\n\n"
    "Syfte och målgrupp påverkade mina val ganska mycket. Eftersom jag vänder mig till kreativa "
    "och estetikmedvetna personer ville jag ha en stilren design med tydlig typografi snarare "
    "än något rörigt eller överfyllt."
)
spacer(doc1)

heading(doc1, "2. Grafiska element", level=2)

heading(doc1, "Typografi", level=3)
body(doc1,
    "Jag valde Playfair Display som rubrikstypsnitt för mitt namn eftersom det är ett elegant "
    "seriftypsnitt som utstrålar kreativitet och professionalism på samma gång. För kontaktuppgifter "
    "och titel använde jag Lato, ett enkelt sans-seriftypsnitt som är enkelt att läsa i liten storlek.\n\n"
    "Hierarkin är uppbyggd så att mitt namn är störst och syns direkt, sedan titeln i lite mindre "
    "storlek, och kontaktinfon i det minsta teckensnittet längst ner. Det gör att ögat naturligt "
    "rör sig i rätt ordning på kortet."
)

heading(doc1, "Färg", level=3)
body(doc1,
    "Jag valde en dämpad rosébeige som bakgrundsfärg på framsidan och svart/mörkgrå text. "
    "Baksidan hade en mörkare sandfärg med ljus text. Färgkombinationen är väldigt neutral men "
    "ändå lite varm och personlig, vilket jag tyckte passade bra för en fotograf.\n\n"
    "Kontrasten mellan bakgrunden och texten är tillräckligt tydlig för att det ska vara läsbart, "
    "men det är inte så skarpt som svart-på-vitt, vilket ger en mjukare och mer konstnärlig känsla."
)

heading(doc1, "Logotyp och bild", level=3)
body(doc1,
    "Logotypen är en enkel stiliserad bokstav 'A' som jag formgav i vecka 4-uppgiften. "
    "Den är gjord i samma mörka färg som brödtexten och placerad i det övre vänstra hörnet på "
    "framsidan. Jag valde att hålla loggan minimal för att den inte ska ta för mycket fokus "
    "från kontaktinfon, men den ger ändå ett tydligt varumärkesintryck."
)

heading(doc1, "Layout och komposition", level=3)
body(doc1,
    "Layouten är ganska luftig. Jag har medvetet lämnat lite tomrum runt textelementen för att "
    "det ska kännas öppet och inte trångt. Texten är vänsterställd på framsidan, vilket ger "
    "ett modernt och avskalat utseende. Baksidan har loggan centrerad med en kort tagline under.\n\n"
    "Balansen är assymetrisk men ändå stabil, loggan och namnet drar ögat på ett naturligt sätt."
)
spacer(doc1)

heading(doc1, "3. Papper och material", level=2)
body(doc1,
    "Jag har valt ett bestruket glansigt papper med en ytvikt på ungefär 350 g/m². Det är ett "
    "ganska tjockt papper vilket gör att kortet känns solitt och kvalitativt när man håller i det.\n\n"
    "Glansen på ytan gör att färgerna ser lite klarare och mer levande ut, vilket passar bra "
    "eftersom jag är fotograf och vill att kortet ska kännas visuellt tilltalande. "
    "Pappret kommunicerar att det finns en viss seriösitet och att jag satsar på kvalité.\n\n"
    "Om jag hade valt ett enkelt obestruket papper hade det känt mer amatörmässigt ut, "
    "vilket inte stämmer med det intryck jag vill ge."
)
spacer(doc1)

heading(doc1, "4. Helhetsbedömning och reflektion", level=2)
body(doc1,
    "Jag tycker att kortet uppfyller sitt syfte ganska väl. Det är tydligt, snyggt och ger "
    "ett professionellt intryck utan att vara överdrivet. Färgvalen och typsnittet hänger ihop "
    "med tanken om en kreativ men seriös fotograf.\n\n"
    "Det som fungerar bäst tycker jag är typografin och den luftiga layouten. Det ser städat ut.\n\n"
    "Om jag fick göra om det skulle jag nog experimentera lite mer med logotypen, "
    "jag känner att den kunde vara lite mer unik och minnesvärdt. Dessutom kanske jag "
    "skulle ha testat ett matt papper istället för glans, för att se om det hade passat "
    "bättre med den varma färgpaletten. Men i helhet är jag nöjd med resultatet."
)

doc1.save(os.path.join(OUT, "INLAMNING_visitkort_analys.docx"))
print("Saved: INLAMNING_visitkort_analys.docx")


# ─────────────────────────────────────────────
# TASK 2 — FOLDER ANALYS (Studiero)
# ─────────────────────────────────────────────
doc2 = Document()

heading(doc2, "Analys och motivering – Folder om studiero", level=1)
body(doc2, "Alyaa\nGrafisk kommunikation – Produktionsprocessen")
spacer(doc2)

heading(doc2, "1. Syfte och målgrupp", level=2)
body(doc2,
    "Syftet med foldern är att informera och uppmuntra elever på Zetterbergsgymnasiet att bidra "
    "till en bättre studiero i skolans gemensamma utrymmen och klassrum. Budskapet ska vara "
    "enkelt och tydligt – vi alla har ett ansvar för hur stämningen i skolan är.\n\n"
    "Målgruppen är eleverna på skolan, alltså ungdomar i åldern 16–19 år. Dom är vana vid digital "
    "kommunikation och reagerar bäst på kortfattad, visuellt tilltalande information snarare "
    "än långa textblock.\n\n"
    "Det här påverkade våra designval ganska direkt. Vi ville ha ett ungdomligt och tillgängligt "
    "formspråk utan att det ser barnsligt ut. Texten är kort och konkret, och illustrationen "
    "är anpassad för att snabbt fånga uppmärksamheten."
)
spacer(doc2)

heading(doc2, "2. Grafiska element", level=2)

heading(doc2, "Typografi", level=3)
body(doc2,
    "Vi valde Open Sans som huvudtypsnitt för hela foldern. Det är ett rent och lättläst "
    "sans-seriftypsnitt som fungerar bra i alla storlekar. Rubriken är i fet stil och betydligt "
    "större än brödtexten för att skapa tydlig hierarki.\n\n"
    "Underrubriker används i mellanstorlek och hjälper läsaren att orientera sig i texten. "
    "Brödtexten är i 11pt vilket är lite mindre men ändå läsbart i det foldrade formatet."
)

heading(doc2, "Färg", level=3)
body(doc2,
    "Vi använde skolans blå färg som en bas, kombinerat med vitt och en ljus gulaktig "
    "accentfärg. Kontrasten mellan mörk blå och vit text är hög och läsbar.\n\n"
    "Gulaccenten användes på puffarna och faktarutorna för att de ska sticka ut lite extra. "
    "Vi tänkte att blå är en lugn och seriös färg vilket passar med ämnet studiero, "
    "medan guldet adderar lite energi och gör det mer inbjudande att läsa."
)

heading(doc2, "Bild, illustration och logotyp", level=3)
body(doc2,
    "Illustrationen som originalgruppen lämnade visar ett klassrum med elever som sitter och "
    "arbetar i lugn och ro. Vi scannande in den och justerade kontrasten lite i Photoshop "
    "för att den skulle passa in i den digitala produktionen.\n\n"
    "Fotot som originalgruppen tog digitalt föreställer ett stillsamt bibliotekshörn i skolan. "
    "Vi beskärde bilden så att fokus hamnar på miljön och stämningen snarare än på enskilda "
    "personer. Det tillförde en realistisk känsla till foldern.\n\n"
    "Skolans logotyp är placerad i nedre högra hörnet på baksidan av foldern för att markera "
    "att det är ett officiellt material från skolan."
)

heading(doc2, "Layout och komposition", level=3)
body(doc2,
    "Foldern är trefaldig och vi delade upp innehållet i tydliga zoner. Framsidan har rubriken "
    "och en stor illustration för att locka in läsaren. Insidan är uppdelad i tre kolumner "
    "med konkreta tips och regler. Baksidan har skolans logotyp och en kort avslutande text.\n\n"
    "Vi har använt en del luft och tomrum för att det inte ska se överfyllt ut. Balansen "
    "är relativt symmetrisk i inneruppslaget men lekfullare på framsidan."
)
spacer(doc2)

heading(doc2, "3. Papper och material", level=2)
body(doc2,
    "Vi tryckte foldern på ett 170 g/m² obestruket papper som tryckerigruppen valde. "
    "Det är ett mellantjockt papper med lite naturlig känsla i ytan, alltså inte blankt eller "
    "glansigt.\n\n"
    "Det obestrukna pappret ger ett mer seriöst och jordnära intryck, vilket passar bra "
    "för ett informationsmaterial som handlar om skolmiljö. Om vi hade valt blankt papper "
    "hade det kanske sett mer kommersiellt ut, som en reklambroschy, vilket inte var tanken.\n\n"
    "Materialet kommunicerar att foldern är informativ och trovärdig, inte ett erbjudande "
    "utan ett budskap från skolan till eleverna."
)
spacer(doc2)

heading(doc2, "4. Helhetsbedömning och reflektion", level=2)
body(doc2,
    "Jag tycker att foldern uppfyller sitt syfte bra. Budskapet är tydligt och designen "
    "är tillräckligt attraktiv för att en elev faktiskt ska vilja läsa den.\n\n"
    "Arbetsprocessen med att vara tre olika grupper (original, produktion och tryckeri) var "
    "väldigt lärorik. Det är inte alltid lätt att jobba med någon annans material, speciellt "
    "när illustrationen inte var exakt vad vi hade förväntat oss. Vi fick anpassa oss.\n\n"
    "Om jag fick göra om det skulle jag vilja att originalgruppen och produktionsgruppen "
    "hade lite mer kontakt i början för att säkerställa att materialet matchar bättre. "
    "Dessutom tycker jag vi hade kunnat leka mer med typografin på framsidan för att "
    "göra den lite mer iögonfallande. Men överlag är jag stolt over resultatet."
)

doc2.save(os.path.join(OUT, "INLAMNING_folder_analys.docx"))
print("Saved: INLAMNING_folder_analys.docx")


# ─────────────────────────────────────────────
# TASK 3 — A3 COLLAGE TEXT (Reklamfoto)
# ─────────────────────────────────────────────
doc3 = Document()

heading(doc3, "Reklamfoto – bildtyp", level=1)
body(doc3, "Alyaa\nGrafisk kommunikation")
spacer(doc3)

heading(doc3, "Vad är reklamfoto?", level=2)
body(doc3,
    "Reklamfoto är en bildtyp som används för att marknadsföra produkter, tjänster eller "
    "varumärken. Syftet är alltid att påverka betraktaren på något sätt, antingen att köpa "
    "något, skapa ett visst intryck av ett varumärke eller väcka ett speciellt känslomässigt svar.\n\n"
    "Till skillnad från ett dokumentärfoto eller reportagefoto är reklambilden nästan alltid "
    "noggrant planerad och styrad. Ingenting i bilden är en slump – belysning, vinkel, "
    "färgsättning och komposition är alla valda med ett syfte."
)
spacer(doc3)

heading(doc3, "Hur används reklamfoto?", level=2)
body(doc3,
    "Reklamfoto syns överallt i vår vardag. Det används i tidningsannonser, på hemsidor, "
    "i sociala medier, på skyltar, förpackningar, katalogser och i tv-reklam.\n\n"
    "Några vanliga exempel:\n"
    "• Matfoto för restauranger och livsmedelsföretag\n"
    "• Modefoto i tidningar och webbutiker\n"
    "• Produktfoto för elektronik, kläder eller kosmetika\n"
    "• Livsstilsfoto som visar hur en produkt används i vardagen\n\n"
    "Gemensamt för alla dessa är att fotografen medvetet skapar en attraktiv och inbjudande "
    "bild som gör betraktaren intresserad av det som visas."
)
spacer(doc3)

heading(doc3, "Syfte och kommunikation", level=2)
body(doc3,
    "Reklamfotots viktigaste syfte är att sälja, men det handlar lika mycket om att bygga "
    "en känsla och ett förtroende kring ett varumärke. En bra reklaambild berättar en liten "
    "historia på bråkdelen av en sekund.\n\n"
    "Fotografen arbetar ofta nära art directors och marknadsförare för att se till att "
    "bilden matchar varumärkets identitet och målgruppens förväntningar. "
    "Färger, ljussättning och modellernas uttryck är alla genomtänkta verktyg för att "
    "nå fram till rätt känsla."
)

doc3.save(os.path.join(OUT, "INLAMNING_collage_text_reklamfoto.docx"))
print("Saved: INLAMNING_collage_text_reklamfoto.docx")

print("\nAlla inlämningsfiles klara!")
