import easyocr
from docx import Document
import os

reader = easyocr.Reader(['en', 'sv'], gpu=False)

BASE = os.path.dirname(os.path.abspath(__file__))

GROUPS = {
    "group1": [
        os.path.join(BASE, "1 (1).jpeg"),
        os.path.join(BASE, "1 (2).jpeg"),
        os.path.join(BASE, "1 (3).jpeg"),
    ],
    "group2": [
        os.path.join(BASE, "2 (1).jpeg"),
        os.path.join(BASE, "2 (2).jpeg"),
        os.path.join(BASE, "2 (3).jpeg"),
    ],
    "one": [
        os.path.join(BASE, "one.jpeg"),
    ],
}

for group_name, images in GROUPS.items():
    doc = Document()
    doc.add_heading(group_name, level=1)

    for img_path in images:
        img_filename = os.path.basename(img_path)
        print(f"Processing: {img_filename} ...")

        results = reader.readtext(img_path, detail=0, paragraph=True)

        doc.add_heading(img_filename, level=2)
        for line in results:
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    out_path = os.path.join(BASE, f"{group_name}.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

print("\nDone! All groups saved.")
