"""
Skapar STUDY_SUMMARY.docx för alla 4 kurser.
Kör: python create_study_summaries.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(os.path.abspath(__file__))

# ── helpers ──────────────────────────────────────────────
def make_doc():
    d = Document()
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(9)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.space_before = Pt(0)
    for sec in d.sections:
        sec.top_margin = Cm(1.2)
        sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
    return d

def title(d, txt):
    p = d.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)

def h1(d, txt):
    p = d.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.space_before = Pt(8)
    p.space_after = Pt(2)

def h2(d, txt):
    p = d.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 80, 80)
    p.space_before = Pt(5)
    p.space_after = Pt(1)

def txt(d, content, bold=False, color=None, size=None):
    p = d.add_paragraph()
    r = p.add_run(content)
    if bold:
        r.bold = True
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.space_after = Pt(1)
    return p

def bullet(d, content):
    p = d.add_paragraph(style="List Bullet")
    p.text = ""
    r = p.add_run(content)
    r.font.size = Pt(8.5)
    p.space_after = Pt(0)

def box(d, label, content):
    """Green concept box"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0, 100, 0)
    r1.font.size = Pt(8.5)
    r2 = p.add_run(content)
    r2.font.size = Pt(8.5)
    p.space_after = Pt(1)

def line(d):
    p = d.add_paragraph()
    r = p.add_run("─" * 70)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(180, 180, 180)
    p.space_before = Pt(2)
    p.space_after = Pt(2)

def save(d, folder, name):
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, name)
    d.save(path)
    print(f"  Saved: {path}")


# ════════════════════════════════════════════════════════════
# 1. ENGELSKA
# ════════════════════════════════════════════════════════════
def create_engelska():
    d = make_doc()
    title(d, "STUDY SUMMARY — Engelska")
    txt(d, "Mohammad Sami Alsharef — Exam prep — alla koncept från kursen", size=8, color=RGBColor(120,120,120))
    line(d)

    # ── A1: Single Story ──
    h1(d, "A1 · The Danger of a Single Story (Chimamanda Ngozi Adichie)")

    h2(d, "Core Concept — What is a single story?")
    txt(d, "A single story = only hearing ONE narrative about a person/group/place → creates stereotypes and incomplete understanding. It's not that the story is wrong, it's that it's INCOMPLETE.")

    h2(d, "Key Term: Nkali")
    box(d, "NKALI", "Igbo word meaning 'to be greater than another'. Whoever has POWER controls which stories get told. Single stories are created by power imbalances — the powerful tell the story, the powerless are defined by it.")

    h2(d, "Adichie's Personal Examples")
    bullet(d, "As a child in Nigeria, wrote stories with white characters eating apples (from British books) → single story of what literature looks like")
    bullet(d, "Fide (house boy): family only known as 'poor' → shocked to see his brother's beautiful basket → her own single story blinded her")
    bullet(d, "American roommate assumed all of Africa = catastrophe, was shocked Adichie spoke English and knew a stove")
    bullet(d, "Adichie in Mexico: herself fell for single story of Mexicans as 'abject immigrants' → felt ashamed when she saw real life")
    bullet(d, "Professor said her novel wasn't 'authentically African' because characters drove cars → single story = Africans must be poor")
    bullet(d, "Student said 'shame Nigerian men are abusers' based on her novel → she replied: 'I read American Psycho, shame all Americans are serial killers'")

    h2(d, "Rhetorical Techniques")
    bullet(d, "Opens with personal hook: 'I'm a storyteller'")
    bullet(d, "Uses HERSELF as example of being wrong (Mexico) → audience doesn't feel attacked")
    bullet(d, "Historical reference: John Lok (1561) wrote about Africans as 'beasts' → shows single stories have centuries of history")
    bullet(d, "Ending: 'when we reject the single story, we regain a kind of paradise' → hopeful conclusion")

    h2(d, "Exam Concepts")
    box(d, "TAKEAWAY", "Stories are powerful. Listen to MULTIPLE stories. Anyone (including us) can fall into the single story trap. It's about power: who tells the story, and who is defined by it.")

    line(d)

    # ── A2: Brian Little ──
    h1(d, "A2 · Brian Little — Who Are You Really? The Puzzle of Personality")

    h2(d, "The OCEAN Model (Big Five)")
    bullet(d, "O = Openness — how curious/creative vs conventional")
    bullet(d, "C = Conscientiousness — how organized/disciplined")
    bullet(d, "E = Extroversion — how energized by social interaction")
    bullet(d, "A = Agreeableness — how cooperative/empathetic")
    bullet(d, "N = Neuroticism — how emotionally reactive/anxious")

    h2(d, "Three Natures of Personality")
    box(d, "BIOGENIC", "What you're born with — your genetic personality baseline (biology)")
    box(d, "SOCIOGENIC", "Shaped by culture, upbringing, social environment")
    box(d, "IDIOGENIC", "What's uniquely YOU — your personal projects and passions")

    h2(d, "Free Traits Theory (Most Important)")
    txt(d, "We can ACT AGAINST our natural personality when we care deeply about something (a 'core project'). Example: Brian Little is an extreme introvert but performs extroversion every morning in class because he LOVES teaching. This is a FREE TRAIT — acting 'out of character' for something meaningful.")

    h2(d, "Key Examples & Evidence")
    bullet(d, "Elbow-licking: extroverts are more likely to try → shows personality affects behavior in all areas")
    bullet(d, "Caffeine study: introverts perform worse with caffeine, extroverts perform better")
    bullet(d, "Sex frequency statistics: extroverts report higher → personality touches everything")
    bullet(d, "Bathroom hiding: Little hides in bathrooms after overstimulation → introverts need 'restorative niches'")

    h2(d, "Speech Structure & Delivery")
    bullet(d, "Hook: 'There are 47 people in this audience suffering from a psychological condition' → creates curiosity")
    bullet(d, "Three-part thesis: 'each of us is like all other people, like some, and like no other person'")
    bullet(d, "Uses rhetorical questions as transitions: 'Are we just a bunch of traits? No.'")
    bullet(d, "Signposting: 'Let me deal in more detail with...', 'First... Second... Third...'")
    bullet(d, "Self-deprecating humor → builds trust. Every joke serves the argument.")
    bullet(d, "Closing: 'Don't follow me' → funny + reinforces introvert message")

    h2(d, "Exam Key Quote")
    box(d, "QUOTE", "'Are we just a bunch of traits? No, we're not.' — The real you is defined by your CORE PROJECTS, not just your born personality.")

    line(d)

    # ── A3: Accents & Identity ──
    h1(d, "A3 · Accents and Identity — TEDx Informative Speech")

    h2(d, "Three-Point Argument Structure")
    bullet(d, "Point 1: Accents = part of identity → should be celebrated, most English speakers are non-native anyway")
    bullet(d, "Point 2: Accent bias has real consequences → less trust, fewer jobs, perceived as less intelligent")
    bullet(d, "Point 3: Solutions → accent-awareness workshops, focus on clarity not 'neutral' accent, practice patience")

    h2(d, "Key Evidence Used")
    box(d, "British Council", "Learners don't need perfect accent to communicate fluently. Thousands of native accent varieties exist.")
    box(d, "Dr Alice Foucart", "Psycholinguist — research shows foreign accents require more cognitive effort to process → leads to negative snap judgments. Flemish study: natives rated as more intelligent than foreigners.")
    box(d, "Trevor Noah", "'An accent is just somebody speaking your language with the rules of theirs' — accents show skill, not lack of ability")

    h2(d, "CRAAP Method (Source Evaluation)")
    bullet(d, "Currency — Is the source recent/up-to-date?")
    bullet(d, "Relevance — Does it directly support your argument?")
    bullet(d, "Authority — Who wrote it? What are their credentials? Is the publisher credible?")
    bullet(d, "Accuracy — Is it backed by evidence/research? Can you cross-check it?")
    bullet(d, "Purpose — Why was it published? Is there bias? Does the bias matter for YOUR use?")

    h2(d, "Speech Skills to Know")
    bullet(d, "Signposting: 'I will make three points', 'first', 'second', 'lastly'")
    bullet(d, "5-paragraph structure: intro → point 1 → point 2 → point 3 → conclusion")
    bullet(d, "Formal register for European Parliament audience")
    bullet(d, "Combining evidence types: research + statistics + quotes")

    line(d)

    # ── General English Skills ──
    h1(d, "General Skills — English Course")

    h2(d, "Discussion Essay Writing")
    bullet(d, "Pick a side but acknowledge counterarguments")
    bullet(d, "Use linking words: however, furthermore, on the other hand, in contrast")
    bullet(d, "Evidence > opinion. Always back claims with examples or sources.")

    h2(d, "Listening Comprehension")
    bullet(d, "Note the speaker's MAIN MESSAGE first, then details")
    bullet(d, "Watch for rhetorical devices: personal anecdotes, statistics, humor, repetition")
    bullet(d, "Identify STRUCTURE: intro → body (usually 3 parts) → conclusion")

    h2(d, "Key Terms Across All Units")
    bullet(d, "Stereotype — oversimplified fixed idea about a group")
    bullet(d, "Anecdote — short personal story used to illustrate a point")
    bullet(d, "Rhetorical question — question asked for effect, not answered")
    bullet(d, "Signposting — phrases guiding the audience through the speech")
    bullet(d, "Free trait — acting against personality for a meaningful project")
    bullet(d, "Accent bias — unconscious prejudice based on how someone sounds")

    save(d, os.path.join(BASE, "engelska", "BM_PREP"), "STUDY_SUMMARY.docx")


# ════════════════════════════════════════════════════════════
# 2. MATEMATIK 2
# ════════════════════════════════════════════════════════════
def create_matematik():
    d = make_doc()
    title(d, "STUDY SUMMARY — Matematik 2")
    txt(d, "Mohammad Sami Alsharef — Exam prep — alla formler & metoder", size=8, color=RGBColor(120,120,120))
    line(d)

    # ── 1. Statistik ──
    h1(d, "1 · Statistik")

    h2(d, "Medelvärde, Median, Typvärde")
    box(d, "MEDELVÄRDE", "Summera alla tal ÷ antal tal. Ex: (4+7+9+2+8)/5 = 6")
    box(d, "MEDIAN", "Sortera först! Mittersta värdet. Udda antal → mitt. Jämnt antal → medel av de två i mitten.")
    box(d, "TYPVÄRDE", "Det tal som förekommer flest gånger.")
    txt(d, "OBS: I en normalfördelning är medelvärde = median = typvärde (alla samma)!")

    h2(d, "Normalfördelning — 68-95-99,7 regeln")
    txt(d, "Kurvan är klockformad och symmetrisk. Toppen = medelvärdet (μ).")
    box(d, "±1σ", "68% av alla värden ligger inom μ ± 1 standardavvikelse")
    box(d, "±2σ", "95% av alla värden")
    box(d, "±3σ", "99,7% av alla värden")
    txt(d, "Ex: μ=170, σ=5 → 68% ligger mellan 165–175. 95% mellan 160–180.")
    txt(d, "Standardavvikelse (σ): liten σ = data nära medelvärdet. Stor σ = utspridd data.")

    h2(d, "Lådagram (Box Plot)")
    txt(d, "Min ──[Q1 ── Median ── Q3]── Max")
    bullet(d, "Q1 = 25%-gräns, Q3 = 75%-gräns")
    bullet(d, "Lådan = mittersta 50% av data (IQR = Q3 − Q1)")
    bullet(d, "Typfråga: 'Vad visar lådan?' → mittersta 50%")

    line(d)

    # ── 2. Algebra ──
    h1(d, "2 · Algebra")

    h2(d, "Lika termer")
    txt(d, "Bara termer med SAMMA variabel kan slås ihop: 3x + 4x = 7x, men 3x + 5 förblir 3x + 5.")

    h2(d, "Multiplicera parenteser")
    txt(d, "3(x+4) = 3x+12.  −2(x−5) = −2x+10 (minus byter tecken på allt)")

    h2(d, "Kvadreringsregler")
    box(d, "(a+b)²", "= a² + 2ab + b². Ex: (x+3)² = x² + 6x + 9")
    box(d, "(a−b)²", "= a² − 2ab + b². Ex: (x−4)² = x² − 8x + 16")
    txt(d, "⚠ VANLIGT MISSTAG: (x+3)² ≠ x²+9. Glöm ALDRIG mittentermen 2ab!", bold=True, color=RGBColor(180,0,0))

    h2(d, "Konjugatregeln")
    box(d, "(a+b)(a−b)", "= a² − b². Ex: (x+5)(x−5) = x² − 25")

    h2(d, "Potensregler")
    bullet(d, "a⁰ = 1")
    bullet(d, "aᵐ · aⁿ = aᵐ⁺ⁿ (multiplicera → addera exponenter)")
    bullet(d, "aᵐ / aⁿ = aᵐ⁻ⁿ (dividera → subtrahera exponenter)")
    bullet(d, "(aᵐ)ⁿ = aᵐ·ⁿ (potens av potens → multiplicera)")
    bullet(d, "a⁻ⁿ = 1/aⁿ (negativ exponent → bråk)")

    h2(d, "Prioriteringsordning (PPPP)")
    txt(d, "1. Parenteser → 2. Potenser → 3. Punkträkning (· /) → 4. Plusminus (+ −)")

    line(d)

    # ── 3. Ekvationer ──
    h1(d, "3 · Ekvationer")

    h2(d, "Linjär ekvation")
    txt(d, "Mål: få x ensamt. Gör SAMMA sak på båda sidor.")
    txt(d, "Ex: 5x−7=8 → 5x=15 → x=3. Kontroll: 5·3−7 = 8 ✓")

    h2(d, "x² = tal")
    txt(d, "x = ±√tal. ALLTID TVÅ SVAR! Ex: x²=64 → x=8 eller x=−8.")

    h2(d, "PQ-formeln (a=1)")
    txt(d, "x² + px + q = 0 → x = −p/2 ± √((p/2)² − q)")

    h2(d, "Allmänna andragradsformeln (alla a)")
    txt(d, "ax² + bx + c = 0 → x = (−b ± √(b²−4ac)) / (2a)")
    txt(d, "Ex: x²−6x+5=0 → a=1,b=−6,c=5 → x=(6±√16)/2 → x=5 eller x=1")

    h2(d, "Ekvationssystem")
    box(d, "SUBSTITUTION", "Lös ut en variabel i ena ekvationen, sätt in i den andra.")
    box(d, "ADDITION", "Addera ekvationerna så att en variabel försvinner.")

    line(d)

    # ── 4. Funktioner ──
    h1(d, "4 · Funktioner")

    h2(d, "Linjär funktion: y = kx + m")
    bullet(d, "k = lutning (k>0 stiger, k<0 sjunker)")
    bullet(d, "m = skärning med y-axeln")
    bullet(d, "Lutning från 2 punkter: k = (y₂−y₁)/(x₂−x₁)")
    bullet(d, "Parallella linjer → samma k")

    h2(d, "Andragradsfunktion: f(x) = ax² + bx + c")
    box(d, "SYMMETRILINJE", "x = −b/(2a) ← VIKTIGASTE FORMELN!")
    box(d, "EXTREMPUNKT", "1) Räkna x med formeln. 2) Sätt in i f(x) → y. 3) Punkt = (x,y)")
    bullet(d, "a > 0 → parabel öppnar uppåt → MINIMUM")
    bullet(d, "a < 0 → parabel öppnar nedåt → MAXIMUM")
    bullet(d, "Nollställen: lös f(x)=0 med andragradsformeln")

    h2(d, "Tillämpning (raket/boll)")
    txt(d, "h(t) = at² + bt + c → max höjd = extrempunkt. c = starthöjd.")
    txt(d, "Ex: h(t)=−5t²+20t+15 → t=−20/(2·(−5))=2s → h(2)=35m")

    line(d)

    # ── 5. Derivata ──
    h1(d, "5 · Derivata")

    h2(d, "Derivationsregler")
    bullet(d, "f(x) = c → f'(x) = 0 (konstant)")
    bullet(d, "f(x) = xⁿ → f'(x) = n·xⁿ⁻¹ (viktigaste!)")
    bullet(d, "f(x) = kx → f'(x) = k")

    h2(d, "Vad derivatan berättar")
    bullet(d, "f'(x) > 0 → funktionen VÄXANDE (stiger)")
    bullet(d, "f'(x) < 0 → funktionen AVTAGANDE (sjunker)")
    bullet(d, "f'(x) = 0 → EXTREMPUNKT")

    h2(d, "Hitta extrempunkt med derivata")
    txt(d, "1) Derivera → f'(x). 2) Lös f'(x)=0 → x. 3) Sätt in x i f(x) → y.")
    txt(d, "Ex: f(x)=x²−6x+5 → f'(x)=2x−6 → 2x−6=0 → x=3 → f(3)=−4 → (3,−4)")

    line(d)

    # ── 6. Exponentialfunktioner ──
    h1(d, "6 · Exponentialfunktioner")
    txt(d, "f(x) = a · bˣ. a = startvärde, b = förändringsfaktor.")
    bullet(d, "b > 1 → TILLVÄXT. Ex: 1,05ˣ = 5% ökning per steg")
    bullet(d, "0 < b < 1 → FÖRFALL. Ex: 0,9ˣ = 10% minskning per steg")
    bullet(d, "Ökar med 15% → multiplicera med 1,15. Minskar 20% → multiplicera med 0,80")

    line(d)

    # ── BM Tips ──
    h1(d, "Inför BM — Praktiskt")
    bullet(d, "TA MED: papper, penna, miniräknare, formelblad (utskrivet)")
    bullet(d, "TA MED: giltig legitimation (OBLIGATORISKT)")
    bullet(d, "FÅR EJ: lärobok, gamla uträkningar, mobil")
    bullet(d, "Kamera PÅ hela mötet")
    bullet(d, "Visa hur du TÄNKER — lika viktigt som rätt svar")

    h2(d, "Vanliga frågetyper")
    bullet(d, "1. Läs av lådagram → Q1, Q3, IQR")
    bullet(d, "2. Normalfördelning → intervall/andel")
    bullet(d, "3. Medelvärde + median från lista")
    bullet(d, "4. Förenkla algebrauttryck")
    bullet(d, "5. Lösa ekvation steg för steg")
    bullet(d, "6. Symmetrilinje + extrempunkt (f(x)=ax²+bx+c)")
    bullet(d, "7. Verklig tillämpning (raket, boll) → andragradsfunktion")
    bullet(d, "8. Potensregler")
    bullet(d, "9. Lutning + skärningspunkt (räta linjen)")
    bullet(d, "10. Derivera enkel funktion")

    save(d, os.path.join(BASE, "matematik-2", "BM_PREP"), "STUDY_SUMMARY.docx")


# ════════════════════════════════════════════════════════════
# 3. PROGRAMMERING 2
# ════════════════════════════════════════════════════════════
def create_programmering():
    d = make_doc()
    title(d, "STUDY SUMMARY — Programmering 2")
    txt(d, "Mohammad Sami Alsharef — Exam prep — Python-koncept från alla uppgifter", size=8, color=RGBColor(120,120,120))
    line(d)

    # ── Grundläggande Python ──
    h1(d, "1 · Grundläggande Python (Uppgift 1: 100-års beräkning)")

    h2(d, "Variabler & Datatyper")
    bullet(d, "str — text: namn = 'Sami'")
    bullet(d, "int — heltal: alder = 18")
    bullet(d, "float — decimaltal: pris = 99.5")
    bullet(d, "bool — True/False")
    bullet(d, "Konvertering: int('25') → 25, str(100) → '100'")

    h2(d, "Input & Output")
    txt(d, "print('text') — skriver ut. input() — läser från användaren.")
    txt(d, "Kursstil: separata print() + input(), INTE input('prompt:')")

    h2(d, "Global vs Lokal")
    txt(d, "'global namn' — gör att en funktion kan ändra en variabel utanför sig. Används i enklare program.")

    h2(d, "Funktioner")
    txt(d, "def funktionsnamn(): ... Bryter ner koden i delar. main() anropar de andra.")

    h2(d, "While-loop med meny")
    txt(d, "fortsatt = True → while fortsatt == True: → visa meny → val = input() → if/elif/else → val=='3' → fortsatt = False")
    txt(d, "OBS: kursstil använder 'while fortsatt == True', INTE 'while True'.")

    line(d)

    # ── Listor & Dicts ──
    h1(d, "2 · Listor, Dictionaries & Sökning (Uppgift 2: Användarhantering)")

    h2(d, "Listor")
    bullet(d, "Skapas: min_lista = []")
    bullet(d, "Lägg till: min_lista.append(element)")
    bullet(d, "Längd: len(min_lista)")
    bullet(d, "Loopa: for item in min_lista:")

    h2(d, "Dictionary (dict)")
    txt(d, "Nyckel-värde-par: anvandare = {} → anvandare['namn'] = 'Ali'")
    txt(d, "Åtkomst: anvandare['namn'] → 'Ali'")
    txt(d, "Lista av dicts = vanligaste sättet att lagra flera objekt.")

    h2(d, "Sökning")
    txt(d, "Konvertera till .lower() före jämförelse → case-insensitive sökning.")
    txt(d, "Loopa igenom lista → kolla 'if sokord in namn_lower' → samla i hittade-lista.")

    h2(d, "Sortering")
    txt(d, "sorted(lista, key=funktion) — sorterar med egen nyckel.")
    txt(d, "Ex: def hamta_namn(a): return a['namn'].lower() → sorted(lista, key=hamta_namn)")

    h2(d, "Tidiga returner")
    txt(d, "if antal == 0: print('Tom') → return — avbryter funktionen tidigt.")

    line(d)

    # ── OOP ──
    h1(d, "3 · OOP — Klasser & Arv (Uppgift 3: Samis Jackets butik)")

    h2(d, "Klass — Grundstruktur")
    txt(d, "class Product:")
    txt(d, "    def __init__(self, name, price): — konstruktor, körs vid skapande")
    txt(d, "    self.name = name — instansvariabel, unik per objekt")
    txt(d, "    def describe(self): — metod, funktion inuti klassen")

    h2(d, "Arv (Inheritance)")
    txt(d, "class Jacket(Product): — Jacket ärver allt från Product")
    txt(d, "Product.__init__(self, name, price) — anropar förälderns konstruktor")
    txt(d, "Jacket har ALLT som Product + egna saker (t.ex. self.size)")

    h2(d, "Method Overriding (Polymorfism)")
    txt(d, "Jacket definierar sin egen describe() som ERSÄTTER Products version.")
    txt(d, "= samma metodnamn, olika beteende beroende på klass.")

    h2(d, "Praktiskt OOP-mönster (butik)")
    bullet(d, "Product — basklass med name, price")
    bullet(d, "Jacket(Product) — ärver + lägger till size")
    bullet(d, "Customer — separat klass med name, email, tier + get_discount()")
    bullet(d, "Cart — self.items = [] → add(), show(), total()")
    bullet(d, "Rabattberäkning: if tier == 'gold': discount = 10 → saved = total * discount / 100")

    h2(d, "self — Vad det är")
    txt(d, "'self' refererar till det specifika objektet. self.name = det HÄR objektets namn. Första parametern i alla metoder.")

    line(d)

    # ── Kursövergripande ──
    h1(d, "4 · Kursövergripande Koncept")

    h2(d, "Programmering 2 — Fokusområden att kunna")
    bullet(d, "OOP: klasser, objekt, arv, polymorfism, inkapsling")
    bullet(d, "Datastrukturer: listor, dictionaries, (stackar/köer)")
    bullet(d, "Filhantering: läsa/skriva filer, JSON, CSV")
    bullet(d, "Felhantering: try/except, egna exceptions")
    bullet(d, "Moduler: import, skapa egna moduler")

    h2(d, "Vanliga Pythonfunktioner")
    bullet(d, "len() — längd/antal")
    bullet(d, "str(), int(), float() — typkonvertering")
    bullet(d, "sorted() — sorterar, returnerar ny lista")
    bullet(d, ".append() — lägg till i lista")
    bullet(d, ".lower() — gör text till gemener")
    bullet(d, "range() — sekvens av tal för loopar")

    h2(d, "Kodstil-tips (kursspecifikt)")
    bullet(d, "Svenska variabelnamn (utan å/ä/ö): anvandare, ratt_svar")
    bullet(d, "while fortsatt == True (inte while True)")
    bullet(d, "Separata print() + input()")
    bullet(d, "Steg-för-steg beräkningar, inte one-liners")
    bullet(d, "Kommentarer på svenska")

    save(d, os.path.join(BASE, "programmering-2", "BM_PREP"), "STUDY_SUMMARY.docx")


# ════════════════════════════════════════════════════════════
# 4. WEBBSERVERUTVECKLING
# ════════════════════════════════════════════════════════════
def create_webb():
    d = make_doc()
    title(d, "STUDY SUMMARY — Webbserverutveckling")
    txt(d, "Mohammad Sami Alsharef — Exam prep — HTML, CSS, JS koncept", size=8, color=RGBColor(120,120,120))
    line(d)

    # ── HTML ──
    h1(d, "1 · HTML5 Grundstruktur")

    h2(d, "Basstruktur")
    txt(d, "<!DOCTYPE html> → <html lang='sv'> → <head> (meta, title, style) → <body> (synligt innehåll)")

    h2(d, "Viktiga Meta-taggar")
    bullet(d, "<meta charset='UTF-8'> — teckenuppsättning")
    bullet(d, "<meta name='viewport' content='width=device-width, initial-scale=1.0'> — responsiv design")
    bullet(d, "<meta name='description' content='...'> — SEO-beskrivning")

    h2(d, "Semantiska Element")
    bullet(d, "<header> — sidhuvud/navigation")
    bullet(d, "<main> — huvudinnehåll")
    bullet(d, "<section> — tematisk grupp")
    bullet(d, "<article> — fristående innehåll")
    bullet(d, "<aside> — sidoinnehåll (sidebar)")
    bullet(d, "<footer> — sidfot")
    bullet(d, "<nav> — navigation")

    h2(d, "Länkar & Media")
    bullet(d, "<a href='url'> — länk. target='_blank' för ny flik")
    bullet(d, "<img src='bild.jpg' alt='beskrivning'> — bild (alt = tillgänglighet)")
    bullet(d, "SVG inline: <svg viewBox='...'><path d='...'/></svg>")

    line(d)

    # ── CSS ──
    h1(d, "2 · CSS — Modern Styling")

    h2(d, "CSS Custom Properties (Variabler)")
    txt(d, ":root { --gold: #c8a96e; --bg: #111; } → Använd: background: var(--bg);")
    txt(d, "Fördel: ändra ETT ställe → ändras överallt. Perfekt för teman (dark/light).")

    h2(d, "CSS Grid Layout")
    txt(d, "display: grid; grid-template-columns: 1fr 250px; gap: 1.5rem;")
    bullet(d, "fr = fraction, delar upp tillgängligt utrymme")
    bullet(d, "grid-column: 1; — placerar element i specifik kolumn")
    bullet(d, "grid-row: 1 / 3; — sträcker sig över rad 1 till 3")

    h2(d, "Flexbox")
    txt(d, "display: flex; → flex-direction: row/column; gap: 0.8rem;")
    bullet(d, "justify-content: center/space-between — horisontell justering")
    bullet(d, "align-items: center — vertikal justering")

    h2(d, "CSS Nesting (Modern)")
    txt(d, ".funk-kort { & .ikon { font-size: 1.8rem; } } → Nästla selektorer direkt. Renare kod.")

    h2(d, "Container Queries")
    txt(d, "container-type: inline-size; → @container funk (max-width: 500px) { ... }")
    txt(d, "Skillnad mot media queries: container query reagerar på ELEMENTETS storlek, inte hela skärmens.")

    h2(d, "Media Queries — Responsiv Design")
    txt(d, "@media (max-width: 600px) { ... } — regler som aktiveras under 600px.")
    txt(d, "Ex: byter grid till 1 kolumn på mobil: grid-template-columns: 1fr;")

    h2(d, "Transitions & Hover")
    txt(d, "transition: background 0.3s; → mjuk animation vid förändring.")
    txt(d, "element:hover { transform: translateY(-2px); } → lyft-effekt.")

    h2(d, "Box Model")
    txt(d, "box-sizing: border-box; → bredd inkluderar padding + border (standard i moderna sidor).")
    bullet(d, "margin — utrymme UTANFÖR elementet")
    bullet(d, "padding — utrymme INNANFÖR elementet")
    bullet(d, "border — ramen runt elementet")
    bullet(d, "border-radius: 1.2rem; — rundade hörn")

    line(d)

    # ── JavaScript ──
    h1(d, "3 · JavaScript — DOM & Events")

    h2(d, "Hämta element")
    txt(d, "document.getElementById('tema-btn') — hämtar elementet med det id:t.")

    h2(d, "Event Listeners")
    txt(d, "element.addEventListener('click', function() { ... }); — reagerar på klick.")

    h2(d, "Ändra klasser (tema-byte)")
    txt(d, "document.body.classList.add('light') → lägger till CSS-klass.")
    txt(d, "document.body.classList.remove('light') → tar bort CSS-klass.")
    txt(d, "Används för dark/light theme: .light { --bg: #f5f0e8; } ändrar CSS-variablerna.")

    h2(d, "Variabler & Logik")
    bullet(d, "var, let, const — deklarerar variabler")
    bullet(d, "if/else — villkor")
    bullet(d, "! (not) — isLight = !isLight → växlar mellan true/false")
    bullet(d, ".textContent — ändrar texten i ett element")

    line(d)

    # ── Praktiskt ──
    h1(d, "4 · Praktiska Koncept (från Uppgift 1: G&S Appen)")

    h2(d, "Vad du byggde")
    txt(d, "En nedladdningssida för Samis Jackets-appen. Dark/light tema, grid layout med sidebar, App Store + Google Play-knappar, container queries, CSS nesting, responsiv design.")

    h2(d, "Tekniker som användes")
    bullet(d, "CSS Variables för teman (:root + .light-klass)")
    bullet(d, "CSS Grid med sidebar (1fr 250px)")
    bullet(d, "CSS Nesting (& .ikon, & .rubrik)")
    bullet(d, "Container Queries (@container funk)")
    bullet(d, "Media Queries (@media max-width: 600px)")
    bullet(d, "Flexbox för navigation och knapplayout")
    bullet(d, "SVG-ikoner inline (App Store, Google Play)")
    bullet(d, "JavaScript: tema-knapp med classList.add/remove")
    bullet(d, "Tailwind CSS (via CDN) för navigation")

    h2(d, "Responsiv Design-strategi")
    bullet(d, "Desktop: 2 kolumner (main + sidebar)")
    bullet(d, "Mobil: 1 kolumn, sidebar under main, mindre rubrik")
    bullet(d, "Knappar: column på mobil, row på desktop")

    line(d)

    # ── Uppgift 2 (M2) ──
    h1(d, "5 · Uppgift 2 (M2): G&S Mediagalleri + Felfix")

    h2(d, "Vad du byggde (Del 1)")
    txt(d, "Fortsättning på G&S/Samis Jackets-sidan från uppgift 1 — samma dark/gold-tema, samma tema-knapp, samma brand. 9 sidor som visar olika multimedia-tekniker, alla kopplade till butikens innehåll.")

    h2(d, "8 sidor + start")
    bullet(d, "Jacka — <img> med beskrivande alt-text (produktbild)")
    bullet(d, "Jingle — <audio controls> med två format (mp3 + ogg)")
    bullet(d, "Reklam — <video controls> i sidan")
    bullet(d, "Logga — <canvas> med JS som ritar G&S-loggan (fillRect + text)")
    bullet(d, "Butik — body med background-image (CSS), fixed bakgrund")
    bullet(d, "Catwalk — <video autoplay muted loop> i wrapper med position:fixed, z-index:-1")
    bullet(d, "3D-jacka — Three.js: BoxGeometry + TextureLoader + animationsloop")
    bullet(d, "VR-butik — A-Frame: <a-scene>, <a-box animation>, <a-sphere>, <a-plane>, <a-sky>, <a-text>")

    h2(d, "Bonus 1 — Tema-knapp som sparar val")
    txt(d, 'localStorage.setItem("tema","light") + classList.add("light"). Vid sidladdning kollar man localStorage och lägger på .light-klassen om den var sparad. Samma id "tema-btn" som i uppgift 1.')

    h2(d, "Bonus 2 — Animerad CSS-bakgrund")
    txt(d, "linear-gradient(120deg, ...) + background-size: 400% 400% + @keyframes rorelse som flyttar background-position från 0% till 100% och tillbaka. animation: rorelse 12s ease infinite.")

    h2(d, "Canvas — så ritar man en form")
    bullet(d, "var c = document.getElementById('minCanvas')")
    bullet(d, "var ctx = c.getContext('2d')")
    bullet(d, "ctx.fillStyle = '#c8a96e'")
    bullet(d, "ctx.fillRect(x, y, bredd, höjd)")
    bullet(d, "ctx.font = 'bold 28px Arial'; ctx.fillText('G&S', 130, 80)")

    h2(d, "Three.js — minsta scen")
    bullet(d, "Scene → PerspectiveCamera → WebGLRenderer.setSize(...).appendChild()")
    bullet(d, "TextureLoader().load(url) → MeshBasicMaterial({map: textur})")
    bullet(d, "BoxGeometry → Mesh(geometri, material) → scene.add(kub)")
    bullet(d, "requestAnimationFrame-loop som roterar och anropar renderer.render(scene, camera)")

    h2(d, "Bakgrundsvideo-tricket")
    txt(d, "video i en wrapper-div med position:fixed; top:0; left:0; width:100%; height:100%; z-index:-1; overflow:hidden. Sen video { width:100%; height:100%; object-fit:cover }. Glöm inte muted (autoplay kräver det).")

    line(d)

    # ── Uppgift 2 — Del 2 ──
    h1(d, "6 · Uppgift 2 (M2): Del 2 — Felfix")

    h2(d, "Värsta JS-buggen")
    txt(d, "Saknad stängande } i function test() — gav SyntaxError och hela scriptet kraschade. Kolla alltid browser-konsolen när JS inte fungerar.")

    h2(d, "Andra typiska fel att leta efter")
    bullet(d, "<img> utan alt-text → tillgänglighetsfel")
    bullet(d, "<input> utan kopplad <label for='id'> → tillgänglighetsfel")
    bullet(d, "Ingen <!DOCTYPE html>, ingen <meta charset>, ingen viewport")
    bullet(d, "Vag länktext typ 'Klicka här' → dåligt för skärmläsare")
    bullet(d, "Dålig kontrast: lightgray text på #eee bakgrund — oläsbart")
    bullet(d, "font-size: 9px → för litet")
    bullet(d, "var x istället för const/let, missade semikolon")
    bullet(d, "Funktion definierad men aldrig anropad")

    h2(d, "Färgkontrast-regel (WCAG)")
    txt(d, "Text mot bakgrund ska ha ratio ≥ 4.5:1 för normal text. lightgray (#D3D3D3) på #eee = nästan samma färg = fail.")

    save(d, os.path.join(BASE, "webbserverutveckling", "BM_PREP"), "STUDY_SUMMARY.docx")


# ════════════════════════════════════════════════════════════
# RUN ALL
# ════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating study summaries...")
    create_engelska()
    create_matematik()
    create_programmering()
    create_webb()
    print("\nDone! 4 STUDY_SUMMARY.docx files created.")
