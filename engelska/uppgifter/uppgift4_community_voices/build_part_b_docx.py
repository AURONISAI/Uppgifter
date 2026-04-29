"""Build Part B Literary Analysis First Draft .docx - clean, simple student format."""
from docx import Document
from docx.shared import Pt, Cm
from pathlib import Path

SRC = Path(__file__).parent / "submit" / "PartB_LiteraryAnalysis_FirstDraft.md"
OUT = Path(__file__).parent / "FINAL_SUBMISSION" / "Assignment4_PartB_LiteraryAnalysis_FirstDraft_MohammadSamiAlsharef.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15


def add_paragraph(text):
    """Add a paragraph; supports *italic* and **bold** inline."""
    p = doc.add_paragraph()
    # split by ** first (bold), then within each part split by * (italic)
    bold_parts = text.split("**")
    for bi, bchunk in enumerate(bold_parts):
        is_bold = (bi % 2 == 1)
        if not bchunk:
            continue
        italic_parts = bchunk.split("*")
        for ii, ichunk in enumerate(italic_parts):
            if not ichunk:
                continue
            r = p.add_run(ichunk)
            r.bold = is_bold
            r.italic = (ii % 2 == 1)
            r.font.size = Pt(12)
    return p


def add_heading(text, size):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


for raw in SRC.read_text(encoding="utf-8").splitlines():
    line = raw.rstrip()
    if not line:
        continue
    if line.startswith("# "):
        add_heading(line[2:].strip(), 15)
    elif line.startswith("## "):
        add_heading(line[3:].strip(), 13)
    else:
        add_paragraph(line)

doc.save(OUT)
print(f"Wrote: {OUT}")
