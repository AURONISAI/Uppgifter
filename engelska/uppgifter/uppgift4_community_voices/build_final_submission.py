"""Build the FINAL_SUBMISSION folder for Assignment 4 (Community Voices)."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, Cm

ROOT = Path(__file__).parent
SUBMIT = ROOT / "submit"
OUT = ROOT / "FINAL_SUBMISSION"
OUT.mkdir(exist_ok=True)

STUDENT = "Mohammad Sami Alsharef"
COURSE = "Engelska 6"
ASSIGNMENT = "Assignment 4 - Part A: Argumentative Speech (Community Voices)"
TOPIC = "Topic 1: Breaking Down Cultural Stereotypes"
SOCIETY = "United Kingdom"

# ---------- 1. Transcript .docx (the main file teacher reads) ----------
speech_md = (SUBMIT / "assignment4_community_voices_speech.md").read_text(encoding="utf-8")

# Strip the markdown headers and turn into clean paragraphs
lines = []
for raw in speech_md.splitlines():
    s = raw.rstrip()
    if s.startswith("## ") or s.startswith("# ") or s.startswith("---"):
        continue
    lines.append(s)
clean = "\n".join(lines).strip()

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

h = doc.add_heading(ASSIGNMENT, level=1)
doc.add_paragraph(f"Student: {STUDENT}")
doc.add_paragraph(f"Course: {COURSE}")
doc.add_paragraph(f"Topic: {TOPIC}")
doc.add_paragraph(f"Society: {SOCIETY}")
doc.add_paragraph(f"Speech length target: 5-6 minutes")
doc.add_paragraph("")

doc.add_heading("Transcript", level=2)
for para in clean.split("\n\n"):
    p = para.strip()
    if not p:
        continue
    if p.lower().startswith("reference list"):
        doc.add_heading("Reference list", level=2)
        continue
    doc.add_paragraph(p)

docx_path = OUT / "Assignment4_Transcript_MohammadSamiAlsharef.docx"
doc.save(docx_path)

# ---------- 2. Plain .txt copy (backup, easy to paste) ----------
shutil.copy2(SUBMIT / "assignment4_community_voices_transcript.txt",
             OUT / "Assignment4_Transcript_MohammadSamiAlsharef.txt")

# ---------- 3. Short version for the audio recording ----------
shutil.copy2(SUBMIT / "assignment4_short_professional_transcript.txt",
             OUT / "Assignment4_AudioScript_Short_MohammadSamiAlsharef.txt")

# ---------- 4. README explaining what to upload ----------
readme = f"""# Assignment 4 - FINAL SUBMISSION

Student: {STUDENT}
Course: {COURSE}
Assignment: {ASSIGNMENT}
Topic: {TOPIC}
Society: {SOCIETY}

## What to upload to the teacher

The assignment requires TWO things:

1. AUDIO RECORDING (sound only) of the speech (5-6 minutes)
   -> Record yourself reading "Assignment4_Transcript_MohammadSamiAlsharef.docx"
   -> Save the audio file in this folder before uploading
   -> Suggested filename: Assignment4_Audio_MohammadSamiAlsharef.mp3 (or .m4a / .wav)

2. TRANSCRIPT of the speech
   -> File: Assignment4_Transcript_MohammadSamiAlsharef.docx (main, formatted)
   -> Backup: Assignment4_Transcript_MohammadSamiAlsharef.txt (plain text)

## Files in this folder

| File                                                     | Purpose                                                |
|----------------------------------------------------------|--------------------------------------------------------|
| Assignment4_Transcript_MohammadSamiAlsharef.docx         | MAIN transcript to submit (formatted, with references) |
| Assignment4_Transcript_MohammadSamiAlsharef.txt          | Same content, plain text backup                        |
| Assignment4_AudioScript_Short_MohammadSamiAlsharef.txt   | Shorter ~3-4 min version if you prefer a tighter audio |
| README.md                                                | This file                                              |

## Sources used in the speech

1. Lesson 13 - Reading in 4 Stages on Cultural Diversity and Stereotypes
   (Echo 2: Open Letter to Fellow Students; Two Kinds; Jaymee's Extended Break)
2. Lesson 14 - Critical Reading on Structural Inequalities and Their Impact
   (Echo 2: To My Old Master; School Shooters Feel Invisible;
    The Worldwide Phenomenon of Black American Culture)
3. UK Government - Ethnicity facts and figures
   https://www.ethnicity-facts-figures.service.gov.uk/

## Recording tips

- Aim for 5:00 - 5:40 (use the long .docx version)
- Speak slowly, pause briefly between the three points
- Use a quiet room, phone microphone is fine
- Do one full take; small mistakes are normal and human
"""
(OUT / "README.md").write_text(readme, encoding="utf-8")

print("FINAL_SUBMISSION folder ready:")
for f in sorted(OUT.iterdir()):
    print(" -", f.name)
