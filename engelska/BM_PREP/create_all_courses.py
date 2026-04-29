from docx import Document
from docx.shared import Pt, RGBColor, Cm
import os

# ============================================================
# SHARED HELPERS
# ============================================================
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(9)
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.space_before = Pt(0)
    for sec in doc.sections:
        sec.top_margin = Cm(1.2)
        sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
    return doc

def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(2)
    for r in heading.runs:
        r.font.size = Pt(11) if level == 1 else Pt(10)
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

def qa(doc, question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run('Q: ' + question + '\n')
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run('-> ' + answer)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x22, 0x66, 0x22)

def line(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(9)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('CODE: ' + text)
    r.font.size = Pt(8)
    r.font.name = 'Consolas'
    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0xB0)

def title(doc, course_name):
    t = doc.add_heading(course_name + ' - What I Submitted', level=0)
    for r in t.runs:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    t.paragraph_format.space_after = Pt(2)
    line(doc, 'Mohammad Sami  |  Green = your actual submitted work')

def save(doc, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f'Saved: {path}')

BASE = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 1. PROGRAMMERING 2
# ============================================================
d = make_doc()
title(d, 'PROGRAMMERING 2')

h(d, 'Uppgift 1 - Interactive Menu (name + age + 100)', level=1)
qa(d, 'What did you build?',
   'Menu program: (1) Enter name+age, (2) Show what year you turn 100, (3) Exit. Uses while-loop, functions, global variables.')
qa(d, 'Functions you wrote?',
   'registrera() = gets name+age via input. visa_hundra() = calculates year (2026 + (100-age)). visa_meny() = prints menu. main() = runs the loop.')
qa(d, 'How does the 100 calculation work?',
   'ar_kvar = 100 - alder, ar_hundra = 2026 + ar_kvar. Prints: "Name, you will turn 100 in the year XXXX."')
qa(d, 'Error handling?',
   'If user picks invalid menu option -> prints "Fel val!" and loops again. If no name registered yet -> "Registera forst!"')

h(d, 'Uppgift 2 - User Management System', level=1)
qa(d, 'What did you build?',
   'User management: (1) Add user (name/age/email), (2) Search by name or email, (3) Show all sorted alphabetically, (4) Exit.')
qa(d, 'Data structure?',
   'List of dictionaries. Each user = {"namn": x, "alder": x, "epost": x}. Stored in global anvandare_lista.')
qa(d, 'How does search work?',
   'Converts search word + name + email to .lower(). Checks if search word is IN name or email. Returns list of matches.')
qa(d, 'How does sorting work?',
   'sorted(anvandare_lista, key=hamta_namn). hamta_namn() returns namn.lower(). Alphabetical sort by name.')
qa(d, 'Functions?',
   'lagg_till_anvandare(), sok_anvandare(), visa_alla_anvandare(), visa_en_anvandare(), hamta_namn(), visa_meny(), main(). 7 total.')
qa(d, 'TODO you left in code?',
   '"todo later maybe to make it real by connecting to maria database in one.com and use it for our customer database"')

h(d, 'Uppgift 3 - OOP: G&S by Samis Jackets Shop', level=1)
qa(d, 'What did you build?',
   'Jacket shop for Samis Jackets. Customer picks products, adds to cart, checks out with membership discount. 4 classes.')
qa(d, 'Your classes?',
   'Product(name, price) -> base class. Jacket(name, price, size) -> inherits from Product, overrides describe(). Customer(name, email, tier) -> has get_discount(). Cart() -> has items list, add/show/total methods.')
qa(d, 'How does inheritance work?',
   'Jacket inherits from Product. Calls Product.__init__(self, name, price) then adds self.size. Overrides describe() to include size.')
qa(d, 'Discount tiers?',
   'bronze=0%, silver=5%, gold=10%, platinum=15%. get_discount() returns the percentage. Checkout calculates: saved = total * discount / 100.')
qa(d, 'Products in the store?',
   'Black Puffer 799 M, Beige Trenchcoat 1099 L, Grey Hoodie 499 S, Winter Parka 1299 XL.')
qa(d, 'Connection to Samis Jackets?',
   'This IS the Samis Jackets app concept. TODO in code: "connect to maria database in one.com, app is live via vercel."')
qa(d, 'OOP concepts shown?',
   'Classes, __init__ constructors, attributes (self.x), methods, inheritance (Jacket from Product), method overriding (describe), lists of objects.')

save(d, os.path.join(BASE, '..', 'programmering-2', 'BM_PREP', 'MY_SUBMISSIONS_Prep.docx'))

# ============================================================
# 2. MATEMATIK 2
# ============================================================
d = make_doc()
title(d, 'MATEMATIK 2')

h(d, 'Uppgift 1 - Modul 1+2 (4 deluppgifter)', level=1)

qa(d, 'Uppgift 1: Normalfordelning - medelvardet?',
   'Tittar pa var kurvan har sin hogsta punkt (toppen). Normalkurvan ar symmetrisk -> toppen = medelvardet. Svar: mu = 9.')

qa(d, 'Uppgift 2: Los ekvationen 5x - 7 = 8?',
   '5x - 7 = 8 -> +7 bada sidor -> 5x = 15 -> /5 -> x = 3. Kontroll: 5*3 - 7 = 8 CHECK.')

qa(d, 'Uppgift 3: f(x) = x^2 - 6x + 5, symmetrilinje + extrempunkt?',
   'a=1, b=-6, c=5. Symmetrilinje: x = -b/(2a) = 6/2 = 3. Extrempunkt: f(3) = 9-18+5 = -4. Punkt: (3, -4). a>0 -> minimum.')

qa(d, 'Uppgift 4: Raket h(t) = -5t^2 + 20t + 15?',
   'a=-5, b=20. Hogsta punkt: t = -20/(2*-5) = 2 sek. h(2) = -20+40+15 = 35m. Svar: 2 sekunder, 35 meters hojd. c=15 = starthojd.')

qa(d, 'Vilka formler anvande du?',
   'Symmetrilinje: x = -b/(2a). Satte in tillbaka i f(x) for att fa y-vardet. Samma formel for bade uppgift 3 och 4.')

save(d, os.path.join(BASE, '..', 'matematik-2', 'BM_PREP', 'MY_SUBMISSIONS_Prep.docx'))

# ============================================================
# 3. WEBBSERVERUTVECKLING
# ============================================================
d = make_doc()
title(d, 'WEBBSERVERUTVECKLING')
line(d, 'Phone Interview Cheat Sheet')
line(d, 'GREEN = what to SAY  |  BLUE = the actual code in your file')

h(d, 'What I Built', level=1)
qa(d, 'What is this page?',
   'A download page for my jacket business app. It has links to App Store and Google Play, a dark and light theme button, and it works on both phones and computers.')

h(d, 'Theme Switch (Dark/Light)', level=1)
qa(d, 'How did you make the theme?',
   'I saved all my colors as variables at the top of the CSS. When someone clicks the theme button, JavaScript adds a class to the page that swaps all those color values at once. So the whole page changes color from one click.')
code(d, ':root { --gold: #c8a96e; --bg: #111; --card: #1a1a1a; }')
code(d, '.light { --bg: #f5f0e8; --card: #fff; --text: #111; }')

qa(d, 'Why use variables for colors?',
   'So I only need to change the color in one place and it updates everywhere on the page automatically. Instead of changing 20 different things one by one.')
code(d, 'background: var(--bg);  /* this reads the variable */')

h(d, 'Page Layout', level=1)
qa(d, 'How is the page structured?',
   'I used CSS Grid to split the page into two columns. The left side has the main content like the download buttons and feature cards. The right side is a smaller sidebar with news updates. The sidebar stretches the full height.')
code(d, 'main { display: grid; grid-template-columns: 1fr 250px; gap: 1.5rem; }')
code(d, 'aside { grid-column: 2; grid-row: 1 / 3; }  /* sidebar spans 2 rows */')

qa(d, 'What happens on a phone screen?',
   'When the screen gets small, the two columns become one column. The sidebar moves below the main content. The download buttons stack on top of each other instead of side by side.')
code(d, '@media (max-width: 600px) { main { grid-template-columns: 1fr; } }')

h(d, 'Container Queries', level=1)
qa(d, 'What is a container query?',
   'It checks how big a specific section is, not the whole screen. So my feature cards switch from two columns to one column based on how much space that section has.')
code(d, '.funk-wrap { container-type: inline-size; container-name: funk; }')
code(d, '@container funk (max-width: 500px) { .funk-grid { grid-template-columns: 1fr; } }')

qa(d, 'How is that different from a normal media query?',
   'A media query looks at the whole window size. A container query looks at just one element. So the features section can adapt on its own even if the screen is big.')

h(d, 'CSS Nesting', level=1)
qa(d, 'What is CSS nesting?',
   'Instead of writing separate style rules for a parent and its children, I wrote the children styles inside the parent. It keeps everything organized in one place.')
code(d, '.funk-kort { & .ikon { font-size: 1.8rem; } & .rubrik { font-weight: 600; } }')

h(d, 'JavaScript', level=1)
qa(d, 'What does your JavaScript do?',
   'It listens for clicks on the theme button. When clicked, it switches a true/false value. If true, it adds a light class to the page so the light colors activate. If false, it removes it and goes back to dark. It also changes the button text.')
code(d, 'var temaBtn = document.getElementById("tema-btn");')
code(d, 'temaBtn.addEventListener("click", function() { isLight = !isLight; ... })')
code(d, 'document.body.classList.add("light");  // activates light theme')
code(d, 'document.body.classList.remove("light");  // back to dark')

h(d, 'Flexbox', level=1)
qa(d, 'Where did you use flexbox?',
   'The navigation bar at the top uses it to push the logo to the left and the links to the right. The download buttons also use it to sit next to each other on bigger screens.')
code(d, 'nav { display: flex; justify-content: space-between; }  /* logo left, links right */')
code(d, '.knappar { display: flex; flex-direction: column; }  /* stacked on mobile */')

h(d, 'HTML Structure', level=1)
qa(d, 'How did you organize the HTML?',
   'I used meaningful tags like header for the navigation, main for the page content, section for the hero area, article for the features, aside for the sidebar, and footer at the bottom. This helps search engines and screen readers understand the page.')
code(d, '<header> <main> <section class="hero"> <article> <aside> <footer>')

h(d, 'Other Things You Might Be Asked', level=1)
qa(d, 'What does box-sizing border-box do?',
   'It makes sure padding and borders are included inside the width I set. Without it, adding padding makes elements wider than I wanted.')
code(d, '* { box-sizing: border-box; }')

qa(d, 'How do hover effects work?',
   'I added smooth transitions so when you hover over the download buttons, the border changes color and the button lifts up a tiny bit. It happens smoothly, not instantly.')
code(d, '.dl-knapp { transition: border-color 0.2s, transform 0.2s; }')
code(d, '.dl-knapp:hover { border-color: var(--gold); transform: translateY(-2px); }')

qa(d, 'Why Tailwind on the nav?',
   'I used Tailwind only for the navigation bar because it was quick for simple things like background color, spacing and alignment. The rest of the page is all custom CSS I wrote myself.')
code(d, '<script src="https://cdn.tailwindcss.com"></script>  /* only used in nav */')

h(d, 'If They Ask What You Would Improve', level=1)
qa(d, 'Good answer:',
   'I would add a QR code so people can scan and download easily. I would also move the CSS to its own file if the site gets bigger. Maybe add animations on the feature cards and a customer reviews section.')

# ---- Uppgift 2 (M2 - lararbedomd) ----
h(d, 'Uppgift 2 - M2 Mediagalleri + Felfix (Lararbedomd)', level=1)
line(d, 'Same Samis Jackets brand and dark/gold theme as Uppgift 1 - this is a continuation, not a new project.')

qa(d, 'What does Del 1 contain?',
   'A G&S Mediagalleri website with 9 pages (start + 8 demos): Jacka (img), Jingle (audio), Reklam (video), Logga (canvas), Butik (background-image), Catwalk (background-video), 3D-jacka (Three.js), VR-butik (A-Frame). All pages share a sticky nav and the same dark/gold theme as the GS appen page.')

qa(d, 'What are the bonuses?',
   'Theme button (dark default / light) that saves the choice in localStorage so it stays across pages. Animated CSS background on the start page that loops between gold and black using @keyframes.')
code(d, 'localStorage.setItem("tema", "light");  // remembers theme between pages')
code(d, '@keyframes rorelse { 0%{background-position:0% 50%} 50%{100% 50%} 100%{0% 50%} }')

qa(d, 'How is Del 1 connected to Uppgift 1?',
   'Same brand (Samis Jackets / G&S), same color palette (#111 background, #c8a96e gold), same theme button id (tema-btn) with the same dark-default behavior, same JS style (var, ==). It is the next page in the same site, not a new idea.')

qa(d, 'Canvas page - what does it draw?',
   'It draws the G&S logo box - a gold rectangle, a black inner rectangle, the text "G&S" in gold and "by Samis Jackets" underneath. The idea is that we could later draw product badges (NY, REA, VIP) on product images this way.')
code(d, 'ctx.fillStyle = "#c8a96e"; ctx.fillRect(20, 20, 280, 140);')

qa(d, 'Three.js page - what is it for?',
   'A textured cube that rotates - meant as a future 360-degree product view. The texture is loaded with TextureLoader and applied to a BoxGeometry. The animation loop uses requestAnimationFrame to rotate the cube on x and y.')

qa(d, 'A-Frame page - what is in the scene?',
   'A 3D scene with a gold rotating box, a red sphere, a dark plane as floor, a black sky and a gold text "Samis Jackets - VR demo". Idea: a mini VR butik people can walk around in.')

qa(d, 'What was Del 2?',
   'Three buggy files from the teachers GitHub - I had to find all the bugs and fix them with Swedish comments explaining each fix.')

qa(d, 'Worst bug in script.js?',
   'A missing closing bracket } after console.log(x). That made the whole script crash with a SyntaxError - nothing in the file ran at all. I added the } and changed var to const, added semicolons, gave the function a real name (visaTal instead of test) and actually called it.')
code(d, 'function visaTal() { const x = 10; console.log(x); }  // FIXED: added }')

qa(d, 'Worst HTML bugs?',
   'No DOCTYPE, no charset, no viewport, img with no alt and a missing image, input with no label, vague "Klicka har" link, useless title. I fixed all of them and explained each one with a Swedish comment.')

qa(d, 'Worst CSS bugs?',
   'lightgray text on #eee background - basically unreadable, fails accessibility. font-size: 9px way too small. lightblue links also unreadable. No responsive rules at all. I changed text to #222, font to 16px, links to dark blue and added a media query.')

qa(d, 'What did you learn from Del 2?',
   'Always open the browser console when something does not work - a missing } showed up immediately as a SyntaxError. Always put alt on images and label on inputs. Always check color contrast.')

save(d, os.path.join(BASE, '..', 'webbserverutveckling', 'BM_PREP', 'MY_SUBMISSIONS_Prep.docx'))

print('\nAll 3 course prep docs created!')
