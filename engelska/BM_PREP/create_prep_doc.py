from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper
def add_heading_colored(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ===== TITLE =====
title = doc.add_heading('English Assessment Meeting – Quick Prep Sheet', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
p = doc.add_paragraph('Mohammad Sami Alsharef  •  Skim this, don\'t memorize. Talk naturally.')
p.runs[0].italic = True

doc.add_paragraph('─' * 60)

# =========================================================
# 1. DISCUSSION ESSAY (A2)
# =========================================================
add_heading_colored('1 · Discussion Essay (A2) – 5-Paragraph Essay', level=1)

note('Fill in your own essay topic below if different. The questions are about YOUR essay.')

bullet('Topic: ', bold_prefix='Topic: ')
p = doc.add_paragraph('  → Say what your essay was about and WHY you chose it (personal interest, current issue, etc.)')

add_heading_colored('Structure (5-paragraph model)', level=2)
bullet('Intro – hook + thesis statement (your main argument)')
bullet('Body 1 – first argument + evidence/quote')
bullet('Body 2 – second argument or counterpoint + evidence')
bullet('Body 3 – third argument or strongest point + evidence')
bullet('Conclusion – restate thesis, final thought')

add_heading_colored('Sources & Quotes', level=2)
bullet('"The Danger of a Single Story" by Chimamanda Ngozi Adichie – shows how one-sided narratives create stereotypes. Good for arguments about perspective, bias, media representation.', bold_prefix='Course text 1: ')
bullet('"Who are you, really?" by Brian Little – personality is more than traits; we act "out of character" for things we care about. Good for arguments about identity, behavior, social expectations.', bold_prefix='Course text 2: ')
bullet('Mention specific quotes you used and explain what they added.')
bullet('Show you compared viewpoints (e.g. "On one hand… on the other hand…")')

add_heading_colored('Source quality (CRAAP)', level=2)
bullet('TED Talks = credible speakers, peer-reviewed ideas, wide audience')
bullet('If you used outside sources: who wrote it, when, why, any bias?')

doc.add_paragraph('─' * 60)

# =========================================================
# 2. INFORMATIVE SPEECH (A3)
# =========================================================
add_heading_colored('2 · Informative Speech (A3) – Accents & Identity', level=1)

add_heading_colored('Topic & Why It Matters', level=2)
bullet('Topic: Accents and Identity in a multilingual Europe')
bullet('Why: 250M+ people live outside home country; accents affect how people are judged')

add_heading_colored('Three Main Points', level=2)
bullet('Accents should be celebrated as part of identity (British Council: "don\'t need a perfect accent to communicate fluently", thousands of accents exist among native speakers)', bold_prefix='Point 1: ')
bullet('Accent bias has real consequences (Horizon Magazine / Dr Alice Foucart: foreign accents require more cognitive effort → negative judgments; Flemish study — natives rated more intelligent)', bold_prefix='Point 2: ')
bullet('We can promote inclusive communication (Trevor Noah quote: "an accent is just somebody speaking your language with the rules of theirs"; accent-awareness workshops; expose students to diverse speakers)', bold_prefix='Point 3: ')

add_heading_colored('Tone & Delivery', level=2)
bullet('Formal tone — TEDx at European Parliament audience')
bullet('Clear signposting: "I will make three points", "Our first point therefore is…"')
bullet('Transitions between each section with topic sentences')

add_heading_colored('Sources Used', level=2)
bullet('British Council article — learners don\'t need perfect accent; thousands of native accents exist', bold_prefix='Source 1: ')
bullet('Horizon Magazine (EU Commission) — Dr Foucart\'s research on cognitive effort + Flemish experiment', bold_prefix='Source 2: ')
bullet('Trevor Noah quote — "accent is not a measure of intelligence"', bold_prefix='Source 3: ')

add_heading_colored('CRAAP Evaluation (Horizon Magazine – your Part B)', level=2)
bullet('Currency – recent, published by EU Commission, migration is a current topic')
bullet('Relevance – directly supported Point 2 with Flemish experiment data')
bullet('Authority – Dr Alice Foucart, psycholinguist; Horizon Magazine = official EU publication')
bullet('Accuracy – empirical research (Flemish listener experiments); cross-checked with British Council')
bullet('Purpose – communicates EU research to public; slight pro-inclusion angle (acknowledged)')

doc.add_paragraph('─' * 60)

# =========================================================
# 3. NOVEL READING PROGRESS
# =========================================================
add_heading_colored('3 · Novel Reading Progress', level=1)

note('You don\'t need to be finished. Just show you\'re reading and thinking.')

bullet('What theme(s) are you noticing? (e.g. identity, belonging, power, growing up)')
bullet('Which character or situation shows that theme?')
bullet('A quote or moment that stood out — say WHY it struck you')
bullet('Language: easy or hard? Any surprises in style or vocabulary?')
bullet('Give specific examples — page numbers or chapter references help')

doc.add_paragraph('─' * 60)

# =========================================================
# QUICK TIPS
# =========================================================
add_heading_colored('Quick Tips for the Meeting', level=1)
bullet('Do NOT memorize. This is a conversation, not a presentation.')
bullet('Use examples from your essay, speech, readings, and novel.')
bullet('Speak clearly and naturally. Notes are fine, scripts are not.')
bullet('Think out loud — teachers want to see HOW you reason.')
bullet('If you don\'t understand a question, ask them to rephrase it.')
bullet('Connect your answers back to the course material when you can.')

# Save
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, 'BM_Prep_English.docx')
doc.save(out_path)
print(f'Saved to: {out_path}')
