from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

def h(text, level=1, color=0x1A478A):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
    return heading

def bold_line(bold_text, rest=""):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    r.font.size = Pt(10.5)
    if rest:
        r2 = p.add_run(rest)
        r2.font.size = Pt(10.5)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'"{text}"')
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def divider():
    p = doc.add_paragraph('─' * 70)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        r.font.size = Pt(8)

def small_note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('ENGELSKA 6 — TOP 40 THINGS TO KNOW', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
p = doc.add_paragraph()
r = p.add_run('Mohammad Sami Alsharef  •  Read this before the test = pass')
r.italic = True
r.font.size = Pt(10)
small_note('Keyword + short description format. Skim fast, understand the idea.')

divider()

# ============================================================
# SECTION A: THE DANGER OF A SINGLE STORY (Adichie)
# ============================================================
h('A · THE DANGER OF A SINGLE STORY — Chimamanda Ngozi Adichie', level=1)
small_note('Course text 1. TED Talk. Used for essay (A2) and discussions.')

numbered(': When you only hear ONE story about a person/place → stereotypes form. You miss the full picture.', bold_prefix='#1  Single Story')
numbered(': "The problem with stereotypes is not that they are untrue, but that they are INCOMPLETE." — Most important line in the talk.', bold_prefix='#2  Stereotypes = Incomplete')
numbered(': As a child in Nigeria she only read British/American books → wrote stories with white characters eating apples in snow → thought books HAD to be like that.', bold_prefix='#3  Adichie as a child')
numbered(': Discovering writers like Chinua Achebe showed her that "people like me — girls with chocolate skin" could exist in literature. Saved her from a single story of what books are.', bold_prefix='#4  African writers changed her')
numbered(': Family house boy. She only knew "they are poor." Visited his village → his brother made a beautiful basket → shocked. Her poverty-only story blinded her.', bold_prefix='#5  Fide (the basket moment)')
numbered(': Roommate assumed Africa = catastrophe. Was shocked Adichie spoke English, played Mariah Carey, could use a stove. Had pity before even meeting her.', bold_prefix='#6  American roommate')
numbered(': Adichie went to Guadalajara expecting "the abject immigrant." Saw real people, felt ashamed. She ALSO fell into the single story trap → no one is immune.', bold_prefix='#7  Mexico — she admits guilt')
numbered(': Igbo word meaning "to be greater than another." Whoever has POWER controls which stories get told. Power = ability to make YOUR story the definitive one.', bold_prefix='#8  Nkali (power & stories)')
numbered(': Palestinian poet Mourid Barghouti — "if you want to dispossess a people, tell their story and start with secondly." WHERE you start a story changes everything.', bold_prefix='#9  "Start with secondly"')
numbered(': Student said "Nigerian men are abusers" based on her novel character. She replied: "I read American Psycho — shame all Americans are serial killers." We don\'t generalize for powerful countries.', bold_prefix='#10  American Psycho comeback')
numbered(': A professor said her novel wasn\'t "authentically African" because characters drove cars and weren\'t starving. Another single story — Africa must = poverty.', bold_prefix='#11  "Not authentically African"')
numbered(': "Stories can dispossess and malign, but also empower and humanize." Don\'t remove stories — ADD more. "When we reject the single story... we regain a kind of paradise."', bold_prefix='#12  Solution: more stories')

divider()

# ============================================================
# SECTION B: BRIAN LITTLE — PERSONALITY
# ============================================================
h('B · WHO ARE YOU, REALLY? — Brian Little', level=1)
small_note('Course text 2. TED Talk. Used for essay (A2) and analysis (A2 assignment).')

numbered(': Five personality dimensions — Openness, Conscientiousness, Extroversion, Agreeableness, Neuroticism. Standard in psychology.', bold_prefix='#13  OCEAN Model (Big Five)')
numbered(': O + C predict life success. E + A help working with people. N = emotional instability.', bold_prefix='#14  What each trait predicts')
numbered(': (1) Biogenic — born with it (neurophysiology), (2) Sociogenic — culture/society, (3) Idiogenic — what makes YOU uniquely you.', bold_prefix='#15  Three Natures')
numbered(': We can act AGAINST our natural personality when we care about a "core project." E.g. introvert parent fights hospital bureaucracy for sick child.', bold_prefix='#16  Free Traits (MOST IMPORTANT)')
numbered(': He\'s an extreme introvert but acts extroverted every morning for his students because he loves teaching. Living proof of his own theory.', bold_prefix='#17  Little himself = proof')
numbered(': Extroverts need stimulation, stand close, use nicknames ("Charles → Chuck → Chuckles Baby"). Introverts prefer quiet, complex language, distance.', bold_prefix='#18  Extrovert vs Introvert differences')
numbered(': Caffeine works better for extroverts. Introverts + coffee + speeded tasks = worse performance. Not because they\'re less capable.', bold_prefix='#19  Caffeine study')
numbered(': Little describes Michael as "perhaps more assertive than normally called for." Tom says: "He\'s an asshole!" Shows how intro/extroverts communicate differently.', bold_prefix='#20  Tom & Michael story')
numbered(': Hides in bathroom to recover from overstimulation. Someone recognizes him: "Hey, is that Dr. Little?" Perfect illustration of introvert recovery needs.', bold_prefix='#21  Bathroom conclusion')
numbered(': Don\'t ask "what TYPE are you?" Ask "what are your CORE PROJECTS?" That\'s what really defines a person.', bold_prefix='#22  His final message')

divider()

# ============================================================
# SECTION C: ACCENTS & IDENTITY (Speech A3)
# ============================================================
h('C · ACCENTS & IDENTITY — Your TEDx Speech (A3)', level=1)
small_note('Your informative speech. Three main points + three sources + CRAAP evaluation.')

numbered(': Accents are part of who you are. British Council: "learners don\'t need a perfect English accent to communicate fluently." Thousands of native accents exist. "Standard" = minority.', bold_prefix='#23  Point 1 — Celebrate accents')
numbered(': Accent bias is real. Horizon Magazine / Dr Alice Foucart: foreign accents require more cognitive effort → brain judges speaker negatively. Flemish study: natives rated more intelligent.', bold_prefix='#24  Point 2 — Bias has consequences')
numbered(': Promote inclusion. Trevor Noah: "An accent is not a measure of intelligence. An accent is just somebody speaking your language with the rules of theirs." → workshops, diverse speakers, patience.', bold_prefix='#25  Point 3 — Inclusive communication')
numbered(': 250M+ people live outside home country. Accent bias affects hiring, courts, classrooms. Not just a feelings issue — real-world damage.', bold_prefix='#26  Why it matters (stat)')
numbered(': Formal TEDx at European Parliament. Used signposting: "I will make three points", "Our first point therefore is…", "My conclusion here…"', bold_prefix='#27  Tone & structure')

divider()

# ============================================================
# SECTION D: SOURCES & CRAAP
# ============================================================
h('D · SOURCES & CRAAP METHOD', level=1)
small_note('Know how to evaluate ANY source. The CRAAP acronym is key.')

numbered(': Currency, Relevance, Authority, Accuracy, Purpose — five tests to judge if a source is trustworthy.', bold_prefix='#28  CRAAP = 5 letters')
numbered(': Is it recent enough? For accent bias = yes, migration is growing, research is new.', bold_prefix='#29  Currency')
numbered(': Does it actually help your argument? Horizon Magazine data directly supported Point 2 of speech.', bold_prefix='#30  Relevance')
numbered(': Who wrote it? Dr Foucart = psycholinguist. Where published? Horizon Magazine = official EU Commission publication. Not a random blog.', bold_prefix='#31  Authority')
numbered(': Based on experiments (Flemish listener study), not just opinions. Cross-checked with British Council article = same findings.', bold_prefix='#32  Accuracy')
numbered(': Communicates EU research to public. Slight pro-inclusion bias (EU-funded) — you acknowledged this, shows critical thinking.', bold_prefix='#33  Purpose')
numbered(': (1) British Council article — accent diversity is normal, (2) Horizon Magazine / Dr Foucart — accent bias research, (3) Trevor Noah quote — accents ≠ intelligence.', bold_prefix='#34  Your 3 speech sources')

divider()

# ============================================================
# SECTION E: ESSAY WRITING (A2)
# ============================================================
h('E · DISCUSSION ESSAY SKILLS (A2)', level=1)
small_note('5-paragraph model. Know the structure and how to use sources.')

numbered(': Intro (hook + thesis) → Body 1 (argument + evidence) → Body 2 (counter/second point + evidence) → Body 3 (strongest point + evidence) → Conclusion (restate + final thought).', bold_prefix='#35  5-Paragraph Model')
numbered(': Don\'t just drop quotes. Introduce them ("Adichie argues that..."), quote the exact words, then EXPLAIN what it means and why it supports your point.', bold_prefix='#36  How to use quotes')
numbered(': Show BOTH sides. "On one hand... on the other hand..." or "While some argue X, others believe Y." Then take YOUR position with evidence.', bold_prefix='#37  Comparing viewpoints')
numbered(': Adichie on single stories/stereotypes/power. Little on personality/identity/free traits. Both are TED Talks = credible speakers, peer-reviewed ideas.', bold_prefix='#38  Course texts in essay')

divider()

# ============================================================
# SECTION F: SPEECH DELIVERY (A3)
# ============================================================
h('F · SPEECH & DELIVERY SKILLS', level=1)

numbered(': "First... Second... Third...", "Let me now turn to...", "Our first point therefore is...", "In conclusion..." Guide the listener through your structure.', bold_prefix='#39  Signposting phrases')
numbered(': Match tone to audience (formal for Parliament/academic, casual for peers). Use rhetorical questions. Vary pace. Don\'t read — speak to people.', bold_prefix='#40  Adapting delivery')

divider()

# ============================================================
# BONUS: QUICK-FIRE Q&A
# ============================================================
h('BONUS — If They Ask You These Questions', level=1)
small_note('Short answers ready to go. Expand with examples when speaking.')

bold_line('Q: What is the danger of a single story?')
bullet('Only hearing ONE story → incomplete picture → stereotypes → you can\'t see people as full humans.')
quote('The problem with stereotypes is not that they are untrue, but that they are incomplete.')

bold_line('Q: Give an example of a single story.')
bullet('Adichie\'s roommate: Africa = catastrophe. Shocked she spoke English, played Mariah Carey, could use a stove.')
bullet('OR: Adichie herself about Mexicans — only saw "abject immigrants" until she visited Guadalajara.')

bold_line('Q: What are "free traits" (Brian Little)?')
bullet('Acting against your natural personality for something you care about. E.g. introvert Little acts extroverted to teach because he loves his students.')

bold_line('Q: Explain the OCEAN model.')
bullet('O=Openness, C=Conscientiousness, E=Extroversion, A=Agreeableness, N=Neuroticism. Five dimensions of personality. O+C predict success, E+A help social work.')

bold_line('Q: What were your 3 speech points?')
bullet('(1) Celebrate accents as identity, (2) Accent bias has real consequences, (3) Promote inclusive communication.')

bold_line('Q: How did you evaluate your source (CRAAP)?')
bullet('Horizon Magazine: recent (Currency), directly useful data (Relevance), Dr Foucart is a psycholinguist + EU publication (Authority), based on experiments (Accuracy), slight pro-inclusion bias which I acknowledged (Purpose).')

bold_line('Q: What course texts did you use and what did they add?')
bullet('Adichie → showed how limited narratives create stereotypes, the role of power in storytelling.')
bullet('Brian Little → showed identity is complex, we\'re more than our traits, free traits let us act against our nature for things we care about.')

bold_line('Q: How did you structure your essay?')
bullet('5-paragraph model: intro with thesis, 3 body paragraphs each with argument + evidence/quote, conclusion restating thesis with final reflection.')

bold_line('Q: Connection between power and stories?')
bullet('"Nkali" — whoever has power controls the narrative. Start a story at different points → completely different meaning. Power = making YOUR story the definitive one.')

bold_line('Q: What does Trevor Noah\'s quote mean?')
bullet('"An accent is just somebody speaking your language with the rules of theirs" = accents show linguistic skill, not lack of intelligence. Rules of mother tongue naturally transfer.')

divider()

# ============================================================
# FINAL REMINDERS
# ============================================================
h('REMEMBER AT THE TEST', level=1)
bullet('This is a CONVERSATION, not a memorized script. They want to see you THINK.')
bullet('Always give EXAMPLES. Never just say "stereotypes are bad" — say WHY and give a specific moment from the talk/speech.')
bullet('If stuck: connect back to Adichie (single story) or Little (free traits). These two cover almost everything.')
bullet('Use the quotes above. Even one well-placed quote shows you know the material.')
bullet('It\'s OK to say "I think..." or "If I remember correctly..." — sounds natural, not scripted.')
bullet('If you don\'t understand a question, ask them to rephrase. That\'s normal and shows confidence.')

# Save
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TOP_40_English_Prep.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
