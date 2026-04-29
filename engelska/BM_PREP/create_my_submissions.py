from docx import Document
from docx.shared import Pt, RGBColor, Cm
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)

for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

def h(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(2)
    for run in heading.runs:
        run.font.size = Pt(11) if level == 1 else Pt(10)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

def qa(question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run('Q: ' + question + '\n')
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run('-> ' + answer)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x22, 0x66, 0x22)

def line(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(9)

# ====== TITLE ======
t = doc.add_heading('WHAT I SUBMITTED - Quick Prep', level=0)
for r in t.runs:
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
t.paragraph_format.space_after = Pt(2)
line('Mohammad Sami  |  All answers = YOUR exact submitted words  |  Green = what you wrote')

# ====== A1: SINGLE STORY ======
h('A1 - Single Story (Adichie) - Your 16 Answers')

qa('Main message?',
   'Only hearing ONE story -> wrong/incomplete ideas -> stereotypes. We need more than one story.')

qa('How does she hook?',
   'Says "I\'m a storyteller", shares childhood anecdote - wrote stories with white characters in Nigeria. Personal + curious.')

qa('What shift after childhood?',
   '"How impressionable and vulnerable we are." Found African writers (Achebe) -> people like her could exist in books.')

qa('Fide -> roommate connection?',
   'Fide = house boy, only knew "they are poor." Shocked by beautiful basket. Same pattern: roommate\'s single story = "Africa = catastrophe."')

qa('Her conclusion?',
   'Reject single story -> "regain a kind of paradise." More stories = better understanding. Hopeful ending.')

qa('Basket moment?',
   'Couldn\'t imagine poor family making something beautiful. Single story blinds you from seeing any other side.')

qa('History of African storytelling?',
   '"If I hadn\'t grown up in Nigeria I\'d think Africa = landscapes waiting to be saved by white foreigner." John Lok 1561: Africans as "beasts with no houses."')

qa('Power & stories?',
   'Single stories happen because of POWER. Igbo word "nkali" = being greater than another. Whoever has power controls the narrative.')

qa('Why Achebe & Camara Laye?',
   'Shifted her thinking. "People like me - girls with chocolate skin" could be in literature. Saved her from single story of books.')

qa('Mexico trip?',
   'Expected "abject immigrants" (from US media). Saw real people working, laughing. Felt ashamed. She was guilty of same thing she criticises.')

qa('American Psycho?',
   'Student: "shame Nigerian men are abusers." She replied: "I read American Psycho - shame Americans are serial killers." We don\'t generalise for powerful countries.')

qa('Her complex identity?',
   'Happy childhood + love BUT ALSO grandfathers in refugee camps, friend died (no water in fire trucks), jam disappeared under military govt. Both sides = who she is.')

qa('"Not authentically African"?',
   'Professor said characters drove cars, were educated, not starving = "not African." She rejects this. Another single story.')

qa('TV station woman?',
   'A messenger read her novel, told her what to write in sequel. Regular Nigerians DO read - assumption they don\'t = single story.')

qa('Does she blame the audience?',
   'No. Uses HER OWN examples (roommate, Mexico) -> anyone can fall in the trap. Audience reflects without feeling judged.')

qa('Power + language + culture?',
   'Whoever has economic/cultural power shapes the dominant story. West wrote about Africa for centuries -> became "official." Other perspectives left out.')

# ====== A2: BRIAN LITTLE ======
h('A2 - Brian Little Analysis - What You Wrote')

qa('Overall idea?',
   'OCEAN model (Big Five) but REAL point = "free traits." We act against our nature when we care about something. Not just WHAT type, but WHAT you\'re doing & why.')

qa('Structure?',
   'Intro: joke (47 symptoms) + "like all people, like some, like no other person" = structure of whole talk. Body: (1) OCEAN short+light, (2) extroversion+3 natures (biogenic/sociogenic/idiogenic)+caffeine+sex stats, (3) free traits. Conclusion: bathroom story, "Don\'t follow me."')

qa('Signposting?',
   '"Let me deal in more detail with extroversion" / "First...Second...Third" / "Are we just a bunch of traits? No." / callback to intro. Rhetorical Qs as transitions.')

qa('Sources?',
   'Research: OCEAN model, caffeine study, sex stats, Susan Cain "Quiet." Anecdotes: Tom+Michael, bathroom, elbow-licking. You wrote: some stats so humorous hard to know how serious - doesn\'t say where they come from.')

qa('Delivery?',
   'Warm, self-deprecating, funny. Calm body language (ironic - introverted). Every joke serves argument. Says "you" -> personal. Living proof of own theory - introvert performing extroversion.')

# ====== A3a: SPEECH ======
h('A3a - Your Speech: Accents & Identity')

qa('Topic?',
   'Accents and Identity. TEDx at European Parliament. "The way we speak reveals our heritage and shapes how we are perceived."')

qa('Point 1?',
   'Celebrate accents. British Council: "don\'t need perfect accent to communicate fluently." Thousands of native accents. Standard = minority.')

qa('Point 2?',
   'Bias has consequences. Horizon/Dr Foucart: foreign accent -> more cognitive effort -> negative judgments. Flemish study: natives rated more intelligent. 250M affected.')

qa('Point 3?',
   'Promote inclusion. Trevor Noah: "An accent is just somebody speaking your language with the rules of theirs." Workshops, diverse speakers, clarity not "neutral" accent.')

qa('Sources?',
   '(1) British Council - accent diversity. (2) Horizon/Foucart - bias research. (3) Trevor Noah - accents != intelligence.')

qa('Tone?',
   'Formal: "Honourable Members." Signposting: "I will make three points: first...second...third."')

# ====== A3b: CRAAP ======
h('A3b - CRAAP Evaluation (Horizon Magazine)')

qa('Currency?', 'Recent, EU Commission. Migration growing -> recent source relevant. Brain/accent research = new field.')
qa('Relevance?', 'Supported Point 2. Used "cognitive effort" finding + Flemish experiment. Real evidence not opinions.')
qa('Authority?', 'Dr Foucart = psycholinguist. Horizon = EU Commission. Not a blog - official, high standards.')
qa('Accuracy?', 'Experiments (Flemish listeners). Empirical. Cross-checked with British Council - same findings.')
qa('Purpose?', 'EU research for public/policymakers. Possible pro-inclusion bias (EU-funded). You acknowledged it.')

# ====== CONNECTIONS ======
h('If They Ask About Connections')

qa('2 course texts?',
   'Adichie -> limited narratives create stereotypes, power controls stories. Little -> identity is complex, free traits = act against nature for what you care about.')

qa('How did you use quotes?',
   'Speech: "don\'t need perfect accent", "speaking your language with rules of theirs." Adichie: "stereotypes incomplete", "regain paradise", "nkali."')

qa('Compare viewpoints?',
   'Adichie: compared HER OWN bias (Mexico) with others\' (roommate). Speech: accents as positive AND source of bias.')

qa('Novel?',
   'Klara and the Sun (Ishiguro). Submitted Part B first draft. Thesis: a person can\'t be copied; what makes Josie "Josie" lives in the people who love her, not inside her body. PEEL x3: (1) Capaldi\'s plan fails on its own logic - "Nothing inside Josie that\'s beyond the Klaras of this world to continue" p.224. (2) Klara in the Yard - "something very special, but it wasn\'t inside Josie. It was inside those who loved her" p.306. (3) Rick + lifting system - "Whatever happens, you and me, we\'ll always be together" p.173.')

# ====== A4: COMMUNITY VOICES ======
h('A4 - Community Voices (UK Stereotypes) - Parts A & B')

qa('Part A topic?',
   'Argumentative speech, 5-6 min. UK audience. Topic 1: breaking down cultural stereotypes. Refuted "stereotypes are just harmless jokes."')

qa('Part A 3 points?',
   '(1) Stereotypes pull people apart - Lesson 13 Open Letter to Fellow Students. (2) We already know what helps - Lesson 14 critical reading method. (3) Local leaders matter - UK Gov Ethnicity facts and figures.')

qa('Part A sources?',
   'Lesson 13 (Open Letter, Two Kinds, Jaymee\'s Extended Break). Lesson 14 (To My Old Master, School Shooters, Black American Culture). UK Government Ethnicity facts and figures.')

qa('Part B novel + thesis?',
   'Klara and the Sun by Kazuo Ishiguro. Thesis (~65 words): a person can\'t be copied even with perfect tech; what makes Josie unique sits in the people who love her, not inside her body.')

qa('Part B PEEL 1 - Capaldi?',
   'Capaldi believes the replica will work. Quote p.224: "Nothing inside Josie that\'s beyond the Klaras of this world to continue." Plan fails on its own logic - even the Mother doesn\'t really believe she could love a copy.')

qa('Part B PEEL 2 - Klara in the Yard?',
   'Klara herself rejects the copy idea by the end. Quote p.306: "There was something very special, but it wasn\'t inside Josie. It was inside those who loved her." The "specialness" is relational.')

qa('Part B PEEL 3 - Rick?',
   'Lifting system locates worth in genes - same mistake Capaldi makes. Quote p.173: "Whatever happens, you and me, we\'ll always be together." Rick\'s value lives in bonds (Josie, his mum, Klara).')

qa('Part B word count?',
   '619 words total. Thesis 64. Bodies 177 / 183 / 189. Within 500-675. No intro, no conclusion (per brief).')

line('')
line('Don\'t recite. Say "in my analysis I wrote..." and explain naturally. These ARE your words.')

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MY_SUBMISSIONS_Prep.docx')
doc.save(out)
print(f'Saved: {out}')
